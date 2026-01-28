import pandas as pd
import io

def process_survey_data(values_file, labels_file, dataset_name=None):
    """
    Merges Qualtrics values and labels datasets into a single DataFrame.
    
    Args:
        values_file: File-like object or path for the Values CSV.
        labels_file: File-like object or path for the Labels CSV.
        dataset_name: Optional string ('pre' or 'post') to customize headers (e.g. Q22 -> RecordID)
        
    Returns:
        pd.DataFrame: The cleaned and merged DataFrame.
    """
    # 1. Ingest
    # We read without header initially to handle the 3-row header structure of Qualtrics
    df_values = pd.read_csv(values_file, header=None)
    df_labels = pd.read_csv(labels_file, header=None)
    
    # 2. Extract Header Info (Row 1 -> QID, Row 2 -> Question Text)
    # 0-based index: Row 0 is QID (e.g. Q1), Row 1 is Text
    # We only care about columns R (index 17) onwards for the merge
    
    # Verify we have enough columns
    if df_values.shape[1] < 18:
        raise ValueError("Dataset has fewer than 18 columns. Expected Qualtrics format starting data at column R.")

    # Slice the relevant parts (Metadata columns A-Q are index 0-16)
    # The user wants to KEEP leading zeros in data, so we treat as string later.
    # But for now, let's just focus on the structure.
    
    # Extract QIDs and Texts from the Labels file (usually safer, though Values should match)
    qids = df_labels.iloc[0, 17:]
    questions = df_labels.iloc[1, 17:]
    
    # 3. Build new Composite Header
    # Format: "Qx. Question Text"
    new_headers = []
    for i, (qid, question) in enumerate(zip(qids, questions)):
        # Clean up potential NaNs or non-strings if any
        q = str(qid).strip()
        t = str(question).strip()
        
        # STRICT Logic for RecordID: Always Index 0 in this sliced list?
        # WAIT: qids/questions are slices starting from 17. So index 0 of this loop IS Column 17.
        if i == 0:
             new_headers.append("RecordID")
        else:
             new_headers.append(f"{q}. {t}")
    
    # 4. Filter Data Rows
    # Rows 0, 1, 2 are headers/metadata. Data starts at row 3.
    data_values = df_values.iloc[3:, 17:].reset_index(drop=True)
    data_labels = df_labels.iloc[3:, 17:].reset_index(drop=True)
    
    # 5. Merge Value and Label Columns
    # We want: Col 1 Value, Col 1 Label, Col 2 Value, Col 2 Label...
    merged_data = pd.DataFrame()
    
    # Ensure they have same number of columns/rows
    # We'll use the headers count to iterate
    for i, header in enumerate(new_headers):
        # Column from Values
        # Note: We need to ensure we are preserving leading zeros. 
        # Pandas read_csv defaulting to infer types might loose them if we aren't careful.
        # However, since we read with header=None, everything is initially object/mixed.
        # Later we will apply whitespace stripping.
        
        col_val = data_values.iloc[:, i].astype(str).str.strip()
        col_lab = data_labels.iloc[:, i].astype(str).str.strip()
        
        # Replace 'nan' string with empty if it occurred due to conversion
        col_val = col_val.replace('nan', '')
        col_lab = col_lab.replace('nan', '')
        
        if header == "RecordID":
             # Force string type, strip whitespace, handle nan
             col_val = col_val.astype(str).str.strip().replace('nan', '')
        elif i > 0:
            # For non-identifier columns, convert Values to numeric (int)
            # errors='coerce' turns non-numeric to NaN
            col_val = pd.to_numeric(col_val, errors='coerce').astype('Int64')

        # Add to new DF
        # We can use the header for both, or differentiate. 
        # Typically "Value" and "Label" suffixes help, but user asked for:
        # "Record the number form values in one column, then record the label... in the adjacent column"
        # "You'll be left with two columns for each question"
        
        merged_data[f"{header} (Value)"] = col_val
        merged_data[f"{header} (Label)"] = col_lab
        
    # Heuristic Check: Do the Label columns look numeric?
    # We check a few columns in the middle
    num_numeric_labels = 0
    check_cols = [c for c in merged_data.columns if "(Label)" in c][:5] # Check first 5 label cols
    for c in check_cols:
        # Check if column is numeric-like (digits)
        if pd.to_numeric(merged_data[c], errors='coerce').notna().sum() > (len(merged_data) * 0.8):
             num_numeric_labels += 1
             
    if num_numeric_labels > 0:
        print("WARNING: It appears your Label columns contain numeric values. Please check if you uploaded the correct 'Labels' file (Choice Text).")


    return merged_data

from docx import Document
from docx.shared import Pt

def generate_docx_dictionary(df):
    """
    Generates a DOCX Data Dictionary from the merged dataframe.
    
    Args:
        df: The merged pd.DataFrame containing (Value) and (Label) columns.
        
    Returns:
        BytesIO: The DOCX file in memory.
    """
    doc = Document()
    doc.add_heading('Data Dictionary', 0)
    
    # Identify Question Groups
    # We look for columns ending in " (Value)"
    # The corresponding label column should be " (Label)"
    
    # We iterate through columns 
    # The structure is usually: Q1 (Value), Q1 (Label), Q2 (Value), Q2 (Label)...
    
    processed_bases = set()
    
    for col in df.columns:
        if col.endswith(" (Value)"):
            base_name = col[:-len(" (Value)")] # e.g. "Q1. Question Text"
            
            if base_name in processed_bases:
                continue
                
            processed_bases.add(base_name)
            
            label_col = f"{base_name} (Label)"
            if label_col not in df.columns:
                continue # Skip if no matching label column
            
            # --- Add Section for this Question ---
            doc.add_heading(base_name, level=2)
            
            # Extract Unique Pairs
            # Drop duplicates and N/A
            pairs = df[[col, label_col]].drop_duplicates().dropna()
            
            # Sort by Value (try numeric sort if possible)
            try:
                pairs['sort_key'] = pd.to_numeric(pairs[col])
                pairs = pairs.sort_values('sort_key')
            except:
                pairs = pairs.sort_values(col)
            
            # Create Table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Value (SPSS)'
            hdr_cells[1].text = 'Label'
            
            # Data Rows
            for index, row in pairs.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row[col])
                row_cells[1].text = str(row[label_col])
                
            doc.add_paragraph() # Spacer
            
    # Save to BytesIO
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f
